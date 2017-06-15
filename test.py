from docx import Document
import os

cte_use_docx = False

heading_str = 'Heading'
heading_len = len(heading_str) + 1
heading_len_odt = len(heading_str) + 4

def iter_headings(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading'):
            yield paragraph, paragraph.style.name, int(paragraph.style.name[heading_len:])


def iter_headings_odt(paragraphs):
    for paragraph in paragraphs:
        style = paragraph.getAttribute('stylename')
        if style.startswith('Heading'):
            yield str(paragraph), style, int(style[heading_len_odt:])
            

def chapter_number(counter, level):
    ret = str(counter[0]) + "."
    for i in range(1, level):
        ret += str(counter[i]) + "."
        
    return ret
    

if cte_use_docx:
    document = Document('input.docx')
    for p in document.paragraphs:
        print p.text

    heading_counter = [0,0,0,0,0,0,0,0,0,0]

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = 'Capitulo'
    hdr_cells[1].text = 'Titulo'
    hdr_cells[2].text = 'Check'

    for heading,htype,hlevel in iter_headings(document.paragraphs):
        if (heading.text):
            print "Heading: ", htype , heading.text
            row_cells = table.add_row().cells
            heading_counter[ hlevel - 1 ] = heading_counter[ hlevel - 1 ] + 1
            for i in range(hlevel, 10):
                heading_counter[i] = 0
            row_cells[0].text = chapter_number(heading_counter,hlevel)
            row_cells[1].text = heading.text
            row_cells[2].text = "YES / NO"  


    # document.add_page_break()
    # document.add_page_break()
    document.save('result.docx')

else:
    from odf.opendocument import OpenDocumentText
    from odf.opendocument import load
    from odf import text
    from odf.table import Table, TableColumn, TableRow, TableCell
    from odf.style import Style, TextProperties, ParagraphProperties
    from odf.style import TableColumnProperties
    from odf.text import P

    doc = load('input.odt')
    
    for paragraph in doc.getElementsByType(text.P):
        print paragraph.getAttribute('stylename')

    heading_counter = [0,0,0,0,0,0,0,0,0,0]

    # Create a style for the table content. One we can modify
    # later in the word processor.
    tablecontents = Style(name="Table Contents", family="paragraph")
    tablecontents.addElement(ParagraphProperties(numberlines="false", linenumber="0"))
    doc.styles.addElement(tablecontents)

    # Create automatic styles for the column widths.
    # We want two different widths, one in inches, the other one in metric.
    # ODF Standard section 15.9.1
    widthshort = Style(name="Wshort", family="table-column")
    widthshort.addElement(TableColumnProperties(columnwidth="1.7cm"))
    doc.automaticstyles.addElement(widthshort)

    widthwide = Style(name="Wwide", family="table-column")
    widthwide.addElement(TableColumnProperties(columnwidth="1.5in"))
    doc.automaticstyles.addElement(widthwide)

    # Start the table, and describe the columns
    table = Table()
    table.addElement(TableColumn(numbercolumnsrepeated=1,stylename=widthshort))
    table.addElement(TableColumn(numbercolumnsrepeated=1,stylename=widthwide))
    table.addElement(TableColumn(numbercolumnsrepeated=1,stylename=widthshort))

    tr = TableRow()
    table.addElement(tr)
    tc = TableCell()
    tr.addElement(tc)
    p = P(stylename=tablecontents,text="Capitulo")
    tc.addElement(p)
    tc2 = TableCell()
    tr.addElement(tc2)
    p2 = P(stylename=tablecontents,text="Titulo")
    tc2.addElement(p2)
    tc3 = TableCell()
    tr.addElement(tc3)
    p3 = P(stylename=tablecontents,text="Check")
    tc3.addElement(p3)
    
    for heading,htype,hlevel in iter_headings_odt(doc.getElementsByType(text.P)):
        if (len(heading)>0):
            print "Heading: ", htype , heading
            heading_counter[ hlevel - 1 ] = heading_counter[ hlevel - 1 ] + 1
            for i in range(hlevel, 10):
                heading_counter[i] = 0

            tr = TableRow()
            table.addElement(tr)
            tc = TableCell()
            tr.addElement(tc)
            p = P(stylename=tablecontents,text=chapter_number(heading_counter,hlevel))
            tc.addElement(p)
            tc2 = TableCell()
            tr.addElement(tc2)
            p2 = P(stylename=tablecontents,text=heading)
            tc2.addElement(p2)
            tc3 = TableCell()
            tr.addElement(tc3)
            p3 = P(stylename=tablecontents,text="YES / NO")
            tc3.addElement(p3)
    
    doc.text.addElement(table)    
    doc.save("result.odt")